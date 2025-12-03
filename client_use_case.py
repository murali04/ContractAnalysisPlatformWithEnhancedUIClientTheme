import os
import io
import re
import json
import uuid
import shutil
import fitz
import spacy
import logging
import urllib.request
import pandas as pd
import numpy as np
import streamlit as st
from datetime import datetime
from dotenv import load_dotenv
from keybert import KeyBERT
from openai import OpenAI
from langdetect import detect
from googletrans import Translator
from langchain_openai import OpenAIEmbeddings
from langchain.vectorstores import FAISS
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.docstore.document import Document
from sklearn.metrics.pairwise import cosine_similarity
# ---------------- Setup ----------------
logging.basicConfig(level=logging.INFO)
load_dotenv()
client = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
embedder = OpenAIEmbeddings(model="text-embedding-3-small")
nlp = spacy.load("en_core_web_sm")
kw_model = KeyBERT()
translator = Translator()
#st.set_page_config(page_title="AI Contract Health Analyzer", layout="wide")
# st.markdown("<h1 style='text-align:center;color:#1F4E79;'>📜 AI-Powered Multilingual Contract Intelligence Dashboard</h1>", unsafe_allow_html=True)
# USER_DIR = "user_memory"
# os.makedirs(USER_DIR, exist_ok=True)
# ---------------- UI / Header Styling ----------------
st.set_page_config(page_title="AI Contract Health Analyzer", layout="wide")
# --- Inject custom professional CSS ---
st.markdown("""
    <style>
        /* Gradient background header */
        .main-header {
            background: linear-gradient(90deg, #0a2540 0%, #144272 35%, #205295 100%);
            padding: 30px 10px;
            border-radius: 12px;
            box-shadow: 0px 4px 10px rgba(0,0,0,0.25);
            text-align: center;
            margin-bottom: 25px;
        }
        .main-header h1 {
            color: white;
            font-size: 38px;
            font-weight: 800;
            letter-spacing: 0.5px;
            text-shadow: 1px 1px 3px rgba(0,0,0,0.3);
            margin-bottom: 8px;
        }
        .main-header p {
            color: #f0f0f0;
            font-size: 18px;
            font-weight: 400;
            letter-spacing: 0.3px;
            margin-top: 0px;
        }
        /* Divider styling */
        hr.divider {
            border: 0;
            height: 2px;
            background: linear-gradient(90deg, rgba(255,255,255,0) 0%, #FFD700 50%, rgba(255,255,255,0) 100%);
            margin: 25px 0;
        }
        /* Info bar styling */
        .info-bar {
            background-color: #f5f7fa;
            border: 1px solid #d0d8e8;
            border-radius: 10px;
            padding: 10px 20px;
            margin-bottom: 25px;
            font-size: 15px;
            font-weight: 500;
            color: #205295;
        }
        /* Metric card hover effect */
        [data-testid="stMetric"] {
            border-radius: 10px;
            background: rgba(255,255,255,0.7);
            padding: 10px;
            box-shadow: 0px 2px 5px rgba(0,0,0,0.05);
            transition: all 0.3s ease;
        }
        [data-testid="stMetric"]:hover {
            transform: scale(1.03);
            background: rgba(255,255,255,0.9);
            box-shadow: 0px 4px 10px rgba(0,0,0,0.1);
        }
    </style>
""", unsafe_allow_html=True)
# --- Header Layout ---
st.markdown("""
    <div class="main-header">
        <h1>🤖 AI-Powered Multilingual Contract Intelligence Dashboard</h1>
        <p>Bringing automation, accuracy, and insight to your enterprise agreements.</p>
        <hr class="divider">
    </div>
""", unsafe_allow_html=True)
# --- Info Bar ---
st.markdown("""
    <div class="info-bar">
        🌐 Supports multilingual contract analysis | 💡 NLP + Generative AI powered reasoning | 🔐 Secure offline embeddings with OpenAI
    </div>
""", unsafe_allow_html=True)
# Create user directory (keep this after header)
USER_DIR = "user_memory"
os.makedirs(USER_DIR, exist_ok=True)
# ---------------- Language Utilities ----------------
def detect_language(text):
    try:
        lang_code = detect(text)
        flags = {
            "en": "🇬🇧", "es": "🇪🇸", "fr": "🇫🇷", "de": "🇩🇪",
            "it": "🇮🇹", "pt": "🇵🇹", "hi": "🇮🇳", "unknown": "🌐"
        }
        return flags.get(lang_code, "🌐")
    except Exception:
        return "🌐"
def translate_to_english(text):
    try:
        lang_code = detect(text)
        if lang_code != "en" and lang_code != "unknown":
            try:
                translated = translator.translate(text, src=lang_code, dest="en")
                return translated.text
            except Exception:
                # fallback to OpenAI translate if googletrans fails
                resp = client.chat.completions.create(
                    model="gpt-4o-mini",
                    messages=[
                        {"role": "system", "content": "Translate the following text to English precisely."},
                        {"role": "user", "content": text}
                    ],
                    temperature=0
                )
                return resp.choices[0].message.content.strip()
        return text
    except Exception:
        return text
# ---------------- Helpers ----------------
def get_user_vector_path():
    session_id = st.session_state.get("session_id", str(np.random.randint(100000, 999999)))
    st.session_state["session_id"] = session_id
    run_id = str(uuid.uuid4())[:8]
    return os.path.join(USER_DIR, f"faiss_{session_id}_{run_id}")
def extract_text_from_pdf(uploaded_file):
    pdf = fitz.open(stream=uploaded_file.read(), filetype="pdf")
    records = []
    for page_num, page in enumerate(pdf, start=1):
        lines = page.get_text("text").split("\n")
        for line_num, line in enumerate(lines, start=1):
            if line.strip():
                records.append({"page": page_num, "line": line_num, "text": line.strip()})
    pdf.close()
    return records
def extract_text_from_docx(uploaded_file):
    from docx import Document as DocxDocument
    doc = DocxDocument(uploaded_file)
    return [{"page": 1, "line": i + 1, "text": p.text.strip()} for i, p in enumerate(doc.paragraphs) if p.text.strip()]
def extract_text_from_excel(uploaded_file):
    df = pd.read_excel(uploaded_file).dropna(how="all")
    df.columns = [str(c).strip() for c in df.columns]
    lines = [" | ".join(map(str, row)) for _, row in df.iterrows()]
    return [{"page": 1, "line": i + 1, "text": t} for i, t in enumerate(lines)]
def chunk_text(records):
    splitter = RecursiveCharacterTextSplitter(chunk_size=2000, chunk_overlap=200)
    docs = []
    for rec in records:
        for idx, chunk in enumerate(splitter.split_text(rec["text"])):
            docs.append(Document(page_content=chunk, metadata={"page": rec["page"], "line": rec["line"], "chunk_id": idx}))
    return docs
def build_vector_store(docs, path):
    if os.path.exists(path):
        shutil.rmtree(path)
    vs = FAISS.from_documents(docs, embedder)
    vs.save_local(path)
    return vs
# ---------------- Keyword generation ----------------
def generate_dynamic_keywords(obligations):
    keywords_dict = {}
    for ob in obligations:
        try:
            keyphrases = kw_model.extract_keywords(ob, keyphrase_ngram_range=(1, 3), stop_words="english", top_n=5)
            doc = nlp(ob)
            noun_phrases = [chunk.text.lower() for chunk in doc.noun_chunks]
            cleaned = list({kp[0].lower() for kp in keyphrases} | set(noun_phrases))
            keywords_dict[ob] = cleaned
        except Exception:
            keywords_dict[ob] = []
    return keywords_dict
# ---------------- RAG + reasoning ----------------
def query_rag(vs, obligation, auto_keywords, top_k=6):
    retriever = vs.as_retriever(search_type="similarity", search_kwargs={"k": top_k})
    docs = retriever.get_relevant_documents(obligation)
    if not docs:
        return {
            "obligation": obligation, "is_present": "No", "reason": "No relevant clauses retrieved.",
            "similarity_score": 0.0, "keyword_hits": [], "confidence": 0.0,
            "page": None, "line": None, "supporting_clauses": []
        }
    ob_emb = embedder.embed_query(obligation)
    doc_embs = [embedder.embed_query(d.page_content) for d in docs]
    sims = cosine_similarity([ob_emb], doc_embs)[0]
    best_idx = int(np.argmax(sims))
    best_doc = docs[best_idx]
    best_score = float(sims[best_idx])
    obligation_keywords = auto_keywords.get(obligation, [])
    keyword_hits = [kw for kw in obligation_keywords if kw.lower() in best_doc.page_content.lower()]
    keyword_ratio = len(keyword_hits) / max(1, len(obligation_keywords))
    # LLM prompt (strict JSON)
    # Cosine Similarity Only for Confidence
    confidence = round(best_score * 100, 1)
    
    # Enhanced Prompt with Strengthened Guidelines
    prompt = f"""
You are a multilingual contract compliance analyst. Translate internally to English.
Analyze the obligation and the relevant clauses. Return valid JSON ONLY.

Task: Determine if the 'Obligation' is fully present and agreed to in the 'Relevant Clauses'.

Guidelines:
1. **Reasonable Efforts to Achieve Result**: If the Obligation requires a specific result (e.g., "remedy infringement", "fix the product"), and the Contract commits to "reasonable commercial efforts" to achieve that result, this is ACCEPTABLE. Return "Yes". The contract may also offer alternatives "in lieu of" the primary remedy (e.g., securing licenses instead of fixing) as long as these alternatives achieve the same end result (non-infringement, continued use).

2. **Strict Remedy Matching for Guarantees**: If the Obligation demands a GUARANTEE with specific remedies (e.g., "must fix or replace"), and the Contract adds a REFUND/REIMBURSE option that allows termination of use (walking away), this undermines the guarantee. Return "No".

3. **Refund as Escape Clause**: If the Obligation requires continued use/access (implied by "guarantee", "must secure rights", "must replace"), but the Contract allows the vendor to simply refund money and terminate, this is NON-COMPLIANCE. Return "No".

Return JSON:
{{
  "is_present": "Yes" or "No",
  "reason": "short 1-2 sentence rationale",
  "suggestion": "If 'No', provide a specific clause suggestion to add or modify to achieve compliance. If 'Yes', return null."
}}

Obligation:
{obligation}

Relevant Clauses:
{chr(10).join([d.page_content for d in docs])}
"""
    llm_status, llm_reason = "No", ""
    try:
        resp = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a concise contract compliance expert who replies in JSON. You only output Yes or No."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.1,
            max_tokens=400
        )
        res_text = resp.choices[0].message.content.strip()
        if "```json" in res_text:
            res_text = res_text.split("```json")[1].split("```")[0]
        elif "```" in res_text:
            res_text = res_text.split("```")[1]
        parsed = json.loads(res_text)
        llm_status = parsed.get("is_present", "No").strip()
        
        # Normalize to Title Case (Yes/No)
        if llm_status.lower() == "yes":
            llm_status = "Yes"
        elif llm_status.lower() == "no":
            llm_status = "No"
        
        # Enforce strict Yes/No
        if llm_status not in ["Yes", "No"]:
            llm_status = "No"
            
        llm_reason = parsed.get("reason", "").strip()
        llm_suggestion = parsed.get("suggestion", None)
        
        # Fix suggestion logic: null/None for Yes, actual suggestion for No
        if llm_status == "Yes":
            llm_suggestion = None
        elif not llm_suggestion or llm_suggestion == "null":
            llm_suggestion = "Consider adding explicit language to address this obligation."
    except Exception:
        llm_reason = "Reason could not be parsed from model response."
        llm_suggestion = "Could not generate suggestion due to error."
    
    # Final Status based on LLM (Semantic Analysis)
    final_status = llm_status
    
    # Removed strict similarity threshold to allow LLM reasoning (semantic match) to prevail
    # even if cosine similarity is low due to vocabulary differences.

    supporting_clauses = [f"[Page {d.metadata.get('page')} Line {d.metadata.get('line')}] {d.page_content[:250].strip()}" for d in docs]
    
    return {
        "obligation": obligation, "is_present": final_status, "reason": llm_reason,
        "similarity_score": round(best_score, 3), "keyword_hits": keyword_hits,
        "confidence": confidence, "page": best_doc.metadata.get("page"), "line": best_doc.metadata.get("line"),
        "supporting_clauses": supporting_clauses
    }
# ---------------- Persistent storage for metrics/history ----------------
if "history" not in st.session_state:
    st.session_state["history"] = []
if "metrics" not in st.session_state:
    st.session_state["metrics"] = {
        "total": 0, "yes": 0, "partial": 0, "no": 0,
        "avg_conf": 0.0, "avg_sim": 0.0, "runtime": 0, "rag_size": 0,
        "last_updated": "Never"
    }
# Create a placeholder to render metrics so we can update in-place
metrics_placeholder = st.empty()
def render_metrics(metrics):
    # clears and renders metrics in the placeholder
    with metrics_placeholder.container():
        st.markdown("### 📊 Contract Health Overview (Latest Analysis)")
        cols = st.columns(4)
        cols[0].metric("📑 Total Obligations", metrics.get("total", 0))
        cols[1].metric("✅ Compliant", metrics.get("yes", 0))
        cols[2].metric("🟡 Partial", metrics.get("partial", 0))
        cols[3].metric("❌ Non-Compliant", metrics.get("no", 0))
        cols2 = st.columns(4)
        cols2[0].metric("🔍 Avg Similarity", f"{metrics.get('avg_sim',0.0):.2f}")
        cols2[1].metric("📊 Avg Confidence", f"{metrics.get('avg_conf',0.0):.1f}%")
        cols2[2].metric("🕐 Runtime", f"{metrics.get('runtime',0)}s")
        cols2[3].metric("⚙️ RAG Index Size", metrics.get("rag_size", 0))
        st.caption(f"🕓 Last Updated: {metrics.get('last_updated','Never')}")
        st.divider()
# render top metrics initially
render_metrics(st.session_state["metrics"])
# ---------------- UI - Sidebar ----------------
st.sidebar.header("📁 Upload Files")
if st.sidebar.button("🧹 Reset History"):
    st.session_state["history"].clear()
    st.session_state["metrics"] = {"total": 0, "yes": 0, "partial": 0, "no": 0, "avg_conf": 0.0, "avg_sim": 0.0, "runtime": 0, "rag_size": 0, "last_updated": "Never"}
    render_metrics(st.session_state["metrics"])
    st.success("History and metrics reset.")
ob_file = st.sidebar.file_uploader("Upload Obligations (CSV/Excel)", type=["csv", "xlsx"])
contract_file = st.sidebar.file_uploader("Upload Contract Document", type=["pdf", "docx", "xlsx", "txt"])
if not ob_file or not contract_file:
    st.warning("Please upload both obligations and a contract file.")
    st.stop()
# ---------------- Load obligations ----------------
if ob_file.name.endswith(".csv"):
    df_ob = pd.read_csv(ob_file).dropna(how="all")
else:
    df_ob = pd.read_excel(ob_file).dropna(how="all")
df_ob.columns = [str(c).strip() for c in df_ob.columns]
st.info("🌐 Detecting & translating obligations...")
# detect language and translate obligations; drop empty lines
df_ob["Language"] = df_ob.iloc[:, 0].astype(str).apply(lambda x: detect_language(x))
df_ob["Obligation_English"] = df_ob.iloc[:, 0].astype(str).apply(lambda x: translate_to_english(x).strip())
df_ob = df_ob[df_ob["Obligation_English"].str.strip() != ""].reset_index(drop=True)
st.markdown("### 🌍 Detected Obligation Languages & Translations")
st.dataframe(df_ob[[df_ob.columns[0], "Language", "Obligation_English"]], use_container_width=True)
obligations = df_ob["Obligation_English"].tolist()
# ---------------- Extract contract text ----------------
if contract_file.name.endswith(".pdf"):
    records = extract_text_from_pdf(contract_file)
elif contract_file.name.endswith(".docx"):
    records = extract_text_from_docx(contract_file)
elif contract_file.name.endswith(".xlsx"):
    records = extract_text_from_excel(contract_file)
else:
    # txt
    text = contract_file.read().decode("utf-8")
    records = [{"page": 1, "line": i+1, "text": l.strip()} for i, l in enumerate(text.split("\n")) if l.strip()]
st.info("🌐 Translating contract to English (if needed)...")
for rec in records:
    rec["text"] = translate_to_english(rec["text"]).strip()
# ---------------- Run Analysis ----------------
if st.button("🚀 Run Multilingual Contract Analysis"):
    start_time = datetime.now()
    with st.spinner("Building semantic vector index..."):
        docs = chunk_text(records)
        vs = build_vector_store(docs, get_user_vector_path())
    with st.spinner("Extracting dynamic keywords..."):
        auto_keywords = generate_dynamic_keywords(obligations)
    results = []
    progress = st.progress(0)
    for i, ob in enumerate(obligations):
        progress.progress((i + 1) / max(1, len(obligations)))
        results.append(query_rag(vs, ob, auto_keywords))
    # build results DataFrame and drop empty rows
    df_res = pd.DataFrame(results).dropna(how="all").reset_index(drop=True)
    if not df_res.empty:
        df_res["Compliance Status"] = df_res["is_present"].map({"Yes": "✅ Yes", "Partial": "🟡 Partial", "No": "❌ No"})
    else:
        df_res = pd.DataFrame(columns=["obligation", "Compliance Status", "similarity_score", "confidence", "keyword_hits", "page", "line", "reason"])
    # compute metrics
    total = len(df_res)
    yes = int((df_res["is_present"] == "Yes").sum()) if "is_present" in df_res else 0
    partial = int((df_res["is_present"] == "Partial").sum()) if "is_present" in df_res else 0
    no = int((df_res["is_present"] == "No").sum()) if "is_present" in df_res else 0
    avg_conf = float(df_res["confidence"].mean()) if "confidence" in df_res and not df_res["confidence"].isnull().all() else 0.0
    avg_sim = float(df_res["similarity_score"].mean()) if "similarity_score" in df_res and not df_res["similarity_score"].isnull().all() else 0.0
    runtime = (datetime.now() - start_time).seconds
    # update persistent metrics and history
    new_metrics = {
        "total": total, "yes": yes, "partial": partial, "no": no,
        "avg_conf": avg_conf, "avg_sim": avg_sim, "runtime": runtime,
        "rag_size": len(docs), "last_updated": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    }
    st.session_state["metrics"] = new_metrics
    st.session_state["history"].append(new_metrics)
    # update top metrics in place (no duplicate below)
    render_metrics(new_metrics)
    st.success("✅ Analysis complete — top metrics updated.")
    # ---------- Compliance report (cleaned) ----------
    st.markdown("### 📋 Multilingual Compliance Report")
    df_clean = df_res[df_res["obligation"].notnull()].copy()
    # Remove any columns that are completely null to avoid showing empty columns
    df_clean = df_clean.dropna(axis=1, how="all")
    if not df_clean.empty:
        st.dataframe(df_clean, use_container_width=True, height=600)
    else:
        st.info("No compliance results to display.")
    # ---------- History ----------
    st.markdown("### 🕓 Previous Analysis History")
    hist_df = pd.DataFrame(st.session_state["history"]).dropna(how="all")
    if not hist_df.empty:
        # drop completely null columns
        hist_df = hist_df.dropna(axis=1, how="all")
        st.dataframe(hist_df, use_container_width=True, height=300)
    else:
        st.info("No previous runs yet.")
    # ---------- Detailed Review ----------
    st.markdown("### 🔍 Detailed Review per Obligation")
    if not df_clean.empty:
        for _, row in df_clean.iterrows():
            with st.expander(f"{row['obligation']}"):
                st.write(f"**Status:** {row.get('Compliance Status','')}")
                st.write(f"**Confidence:** {row.get('confidence','')}% | **Similarity:** {row.get('similarity_score','')}")
                st.write(f"**Page:** {row.get('page','')} | **Line:** {row.get('line','')}")
                st.write(f"**Reason:** {row.get('reason','')}")
                kw = row.get("keyword_hits", [])
                st.write(f"**Keyword Matches:** {', '.join(kw) if isinstance(kw, list) and kw else 'None'}")
                if isinstance(row.get("supporting_clauses"), list) and row.get("supporting_clauses"):
                    st.markdown("**Supporting Clauses:**")
                    for c in row["supporting_clauses"]:
                        st.markdown(f"- {c}")
    else:
        st.info("No detailed obligation rows.")
    # ---------- Download results ----------
    out = io.BytesIO()
    with pd.ExcelWriter(out, engine="xlsxwriter") as writer:
        df_clean.to_excel(writer, index=False, sheet_name="Contract Analysis")
    st.download_button("📥 Download Excel Report", data=out.getvalue(),
                       file_name="multilingual_contract_analysis.xlsx",
                       mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
